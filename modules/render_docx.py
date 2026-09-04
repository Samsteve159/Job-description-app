"""ATS-safe .docx rendering.

Everything here exists to survive a resume parser, which is a much dumber reader than a
human. The rules that matter, in rough order of how often they break parsing:

  no tables            parsers flatten them in unpredictable column order
  no text boxes        content inside them is frequently dropped entirely
  no headers/footers   often skipped, so contact details there vanish
  no floating shapes   the old CV has one, which is why it was failing
  single column        two-column layouts interleave into nonsense
  real list paragraphs not manual bullet glyphs
  standard headings    "EXPERIENCE", not "Where I've Made An Impact"
  plain hyphen dates   "Apr 2018 - Oct 2021", not en dashes

The gating rule from the plan is enforced here rather than in the UI, because the UI is
not the last line of defence. A block that cites no fact, or a reach that has not been
accepted, cannot reach the page even if a caller asks for it.
"""
from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor

from modules import house

BODY_FONT = "Calibri"
BODY_SIZE = Pt(10.5)
NAME_SIZE = Pt(20)
HEADING_SIZE = Pt(11)
INK = RGBColor(0x14, 0x18, 0x1D)
MUTED = RGBColor(0x44, 0x4C, 0x55)

SECTION_ORDER = ["summary", "skills", "experience", "education", "certifications"]
SECTION_HEADINGS = {
    "summary": "SUMMARY",
    "skills": "CORE SKILLS",
    "experience": "PROFESSIONAL EXPERIENCE",
    "education": "EDUCATION",
    "certifications": "CERTIFICATIONS",
}


# Characters a model reaches for that an ATS parser has no reason to handle well. The
# em dash is the famous one, but the first live run leaked U+2011 NON-BREAKING HYPHEN
# into four blocks, and audit() only screened for em and en dashes, so it passed clean.
# Anything non-ASCII in a resume is a liability: parsers built in the 2000s mangle it,
# and a keyword filter comparing "AI\u2011agent" to "ai agent" scores zero.
PUNCTUATION_MAP = {
    "\u2014": ", ",   # em dash
    "\u2013": "-",    # en dash
    "\u2012": "-",    # figure dash
    "\u2011": "-",    # non-breaking hyphen
    "\u2010": "-",    # hyphen
    "\u2018": "'",    # curly quotes
    "\u2019": "'",
    "\u201c": '"',
    "\u201d": '"',
    "\u2026": "...",  # ellipsis
    "\u00a0": " ",    # non-breaking space
    "\u200b": "",     # zero width space
    "\u2022": "",     # bullet glyph. The list style supplies the bullet
    "\u2192": " to ", # arrow
    "\u00ad": "",     # soft hyphen
}


def plain_text(text: str) -> str:
    """Force text down to characters a parser from 2004 would survive.

    Applied at write time rather than trusted to the prompt. The house style tells the
    model not to do this; this makes it true regardless of whether it listened.
    """
    for bad, good in PUNCTUATION_MAP.items():
        text = text.replace(bad, good)
    return " ".join(text.split())


class BlockedContentError(RuntimeError):
    """Raised when a caller tries to render content that has not cleared the gate."""


@dataclass
class Role:
    title: str
    org: str
    dates: str
    location: str = ""
    bullets: List[str] = field(default_factory=list)
    # A study period sits in the experience run as a dated entry rather than being left
    # for the education section to explain from the bottom of page two. A reader who
    # meets the gap and its reason in the same glance never forms the doubt.
    is_study: bool = False


@dataclass
class ResumePayload:
    name: str
    contact: List[str]                      # ["email", "phone", "linkedin", "location"]
    # The posting's own job title, written under the name. Parsers and recruiters both
    # weight a title match, and this is the one place their exact title can appear without
    # any claim about the past attached to it: it says what is being applied for.
    headline: str = ""
    summary: str = ""
    skills: List[str] = field(default_factory=list)
    # [(label, [term, term]), ...]. Three or four labelled rows read better and screen
    # better than one long pipe-separated line, which is where a filter looks hardest.
    skill_groups: List[Any] = field(default_factory=list)
    experience: List[Role] = field(default_factory=list)
    education: List[str] = field(default_factory=list)
    certifications: List[str] = field(default_factory=list)


def export_name(name: str, title: str, package_id: Optional[int] = None,
                kind: str = "Resume") -> str:
    """The filename a recruiter sees, and that some parsers index.

    "package-18.docx" tells the person opening it nothing and tells the parser less.
    Pass package_id for the copy kept on disk, where two applications for the same title
    at different companies must not overwrite each other. Leave it off for the download,
    where the id is noise to the person receiving it.
    """
    def clean(value: str) -> str:
        value = plain_text(value or "")
        value = re.sub(r"[^A-Za-z0-9 ,.&+-]", " ", value)
        return " ".join(value.split()).strip(" .-")[:60]

    parts = [clean(name) or "Resume", clean(title), kind if kind != "Resume" else ""]
    stem = " - ".join(part for part in parts if part)
    suffix = f" ({package_id})" if package_id is not None else ""
    return f"{stem}{suffix}.docx"


# ------------------------------------------------------------------------ the gate

def gate(blocks: List[dict]) -> List[dict]:
    """Return only blocks allowed onto the page. Raises if a blocked block is present.

    blocks are dicts with at least: text, grade, accepted, fact_ids
    """
    allowed: List[dict] = []
    for block in blocks:
        grade = (block.get("grade") or "blocked").lower()
        fact_ids = block.get("fact_ids") or []

        if grade == "blocked" or not fact_ids:
            raise BlockedContentError(
                f"Block cites no facts and cannot be rendered: {block.get('text', '')[:80]!r}"
            )
        if grade in {"inferred", "stretch"} and not block.get("accepted"):
            raise BlockedContentError(
                f"Block graded {grade!r} has not been accepted: {block.get('text', '')[:80]!r}"
            )
        allowed.append(block)
    return allowed


# --------------------------------------------------------------------- docx writing

def _configure(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.font.color.rgb = INK
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.05
    # He asked for justified body text and overrode the usual advice against it. The
    # advice is about ragged word spacing, and he has seen it and made the call, so this
    # is settled and not to be raised again. Headings, role titles, date lines and the
    # header stay left aligned, which is handled at each call site rather than here.
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # A single line stranded across a page break reads as a document nobody looked at.
    pf.widow_control = True

    for section in doc.sections:
        section.top_margin = Pt(40)
        section.bottom_margin = Pt(40)
        section.left_margin = Pt(48)
        section.right_margin = Pt(48)


def _heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = HEADING_SIZE
    run.font.name = BODY_FONT
    run.font.color.rgb = INK


def _line(doc: Document, text: str, bold: bool = False, muted: bool = False,
          size: Optional[Pt] = None, space_before: int = 0,
          justify: bool = False) -> None:
    p = doc.add_paragraph()
    # Only running prose is justified. A role title or a date line justified to the
    # margin is stretched whitespace with nothing gained, which is the failure mode the
    # advice against justification is actually about.
    p.paragraph_format.alignment = (
        WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    )
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(plain_text(text))
    run.bold = bold
    run.font.name = BODY_FONT
    run.font.size = size or BODY_SIZE
    run.font.color.rgb = MUTED if muted else INK


def _bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(plain_text(text))
    run.font.name = BODY_FONT
    run.font.size = BODY_SIZE
    run.font.color.rgb = INK


def render_resume(payload: ResumePayload, out_path: Path) -> Path:
    """Write an ATS-safe .docx. No tables, text boxes, headers, footers or shapes."""
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _configure(doc)

    # name, as ordinary body text rather than a Heading style so parsers read it plainly
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(plain_text(payload.name))
    run.bold = True
    run.font.size = NAME_SIZE
    run.font.name = BODY_FONT
    run.font.color.rgb = INK

    if payload.headline:
        _line(doc, plain_text(payload.headline), bold=True)

    if payload.contact:
        _line(doc, " | ".join(payload.contact), muted=True)

    if payload.summary:
        _heading(doc, SECTION_HEADINGS["summary"])
        _line(doc, payload.summary, justify=True)

    if payload.skill_groups or payload.skills:
        _heading(doc, SECTION_HEADINGS["skills"])
        if payload.skill_groups:
            for label, terms in payload.skill_groups:
                if terms:
                    _line(doc, f"{label}: " + " | ".join(terms), justify=True)
        else:
            _line(doc, " | ".join(payload.skills), justify=True)

    if payload.experience:
        _heading(doc, SECTION_HEADINGS["experience"])
        for role in payload.experience:
            heading = f"{role.title}  |  {role.org}" if role.org else role.title
            _line(doc, heading, bold=True, space_before=8)
            meta = role.dates if not role.location else f"{role.dates}  |  {role.location}"
            _line(doc, meta, muted=True)
            for bullet in role.bullets:
                _bullet(doc, bullet)

    if payload.education:
        _heading(doc, SECTION_HEADINGS["education"])
        for item in payload.education:
            _line(doc, item)

    if payload.certifications:
        _heading(doc, SECTION_HEADINGS["certifications"])
        for item in payload.certifications:
            _line(doc, item)

    # Every engagement on the record is named by descriptor because the client names
    # belong to his employer. Without this line a reader takes "a national QSR chain" for
    # vagueness rather than discretion, and vagueness is the more expensive reading.
    if _names_an_engagement(payload):
        _line(doc, house.WITHHELD_LINE, muted=True, size=Pt(8.5), space_before=10)

    if _runs_past_one_page(payload):
        _page_footer(doc, payload.name)

    doc.save(str(out_path))
    return out_path


# Characters of body text that fit on one page at 10.5pt with these margins. Measured off
# a rendered two-pager rather than calculated, and it is an estimate: python-docx does not
# lay out pages, so nothing here can know the real count. It only decides whether a footer
# is worth adding, and the cost of being wrong either way is one redundant line.
_CHARS_PER_PAGE = 3600


def _runs_past_one_page(payload: "ResumePayload") -> bool:
    body = [payload.summary or ""] + list(payload.skills or [])
    for role in payload.experience or []:
        body.append(role.title or "")
        body.extend(role.bullets or [])
    body.extend(payload.education or [])
    return sum(len(x) for x in body) > _CHARS_PER_PAGE


def _page_footer(doc: Document, name: str) -> None:
    """Name and page number, on documents that run to a second page.

    The rule elsewhere in this file is that nothing goes in a footer, because a parser
    that skips one loses whatever was in it. That rule is about contact details. A name
    and a page number are a safety net for the printed copy: if a parser drops them
    nothing is lost, because both are already in the body. So the ban stays for contact
    details and lifts for these two, and `audit` checks the distinction rather than the
    presence of a footer.
    """
    footer = doc.sections[0].footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(plain_text(name) + "    ")
    run.font.name = BODY_FONT
    run.font.size = Pt(8.5)
    run.font.color.rgb = MUTED

    # A live page number is a field, and python-docx has no wrapper for one, so the three
    # runs Word expects are written by hand.
    field = para.add_run()
    field.font.name = BODY_FONT
    field.font.size = Pt(8.5)
    field.font.color.rgb = MUTED
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, end):
        field._r.append(node)


def _names_an_engagement(payload: "ResumePayload") -> bool:
    """Whether anything in the document stands in for a client name."""
    text = " ".join(
        [payload.summary or ""]
        + [b for role in (payload.experience or []) for b in (role.bullets or [])]
    ).lower()
    return any(marker in text for marker in _ENGAGEMENT_MARKERS)


# The descriptors the facts use in place of a client name. Each is a phrase the writing
# only produces when it is standing in for one.
_ENGAGEMENT_MARKERS = (
    "a national", "an asx-listed", "a top-5", "a top-five", "an australian",
    "a multinational", "a global", "client", "an fmcg", "a listed",
)


# ------------------------------------------------------------------------ ATS audit

# An email, a phone number or a profile URL. The three things that must never be the only
# copy of themselves, and the only three that made the old blanket footer ban worth having.
_CONTACT_IN_CHROME = re.compile(
    r"[\w.+-]+@[\w-]+\.[\w.]+"
    r"|\+?\d[\d\s().-]{7,}\d"
    r"|(?:linkedin\.com|github\.com)/[\w/-]+",
    re.I,
)

HOSTILE_TAGS = {
    "w:tbl": "table",
    "w:drawing": "floating drawing or image",
    "w:pict": "vml picture",
    "w:txbxContent": "text box",
    "mc:AlternateContent": "fallback shape content",
}


@dataclass
class CoverPayload:
    name: str
    contact: List[str]
    greeting: str = "Dear Hiring Manager,"
    paragraphs: List[str] = field(default_factory=list)
    sign_off: str = "Kind regards,"
    role: str = ""
    company: str = ""
    date_line: str = ""


def render_cover(payload: CoverPayload, out_path: Path) -> Path:
    """Write an ATS-safe cover letter.

    Same constraints as the resume, for the same reason: some systems parse the cover
    letter too, and the ones that do not still store it as a file somebody opens. No
    tables, no headers, no text boxes, and every run through plain_text.

    Refuses an empty letter rather than writing a page containing a greeting and a
    signature, which is worse than no file at all because it looks like it worked.
    """
    body = [p for p in payload.paragraphs if (p or "").strip()]
    if not body:
        raise BlockedContentError(
            "a cover letter with no paragraphs is not a cover letter. Accept at least "
            "one, or fix what the truth gate refused"
        )

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _configure(doc)

    p = doc.add_paragraph()
    run = p.add_run(plain_text(payload.name))
    run.bold = True
    run.font.size = Pt(15)
    run.font.name = BODY_FONT
    run.font.color.rgb = INK

    if payload.contact:
        _line(doc, " | ".join(payload.contact), muted=True)
    if payload.date_line:
        _line(doc, payload.date_line, muted=True, space_before=10)

    if payload.role or payload.company:
        subject = "Application: " + " at ".join(
            part for part in (payload.role, payload.company) if part)
        _line(doc, subject, bold=True, space_before=12)

    _line(doc, payload.greeting, space_before=12)
    for text in body:
        _line(doc, text, space_before=8)

    _line(doc, payload.sign_off, space_before=14)
    _line(doc, payload.name)
    doc.save(str(out_path))
    return out_path


def audit(path: Path) -> Dict[str, object]:
    """Re-open a rendered docx and report what a parser would actually see."""
    import zipfile

    path = Path(path)
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", "replace")
        names = zf.namelist()
        # Read inside the with. The first version of this reached for zf further down the
        # function, by which point the archive was closed.
        chrome = "\n".join(
            zf.read(n).decode("utf-8", "replace")
            for n in names
            if n.startswith(("word/header", "word/footer"))
        )

    doc = Document(str(path))
    text = "\n".join(p.text for p in doc.paragraphs)

    problems = []
    for tag, label in HOSTILE_TAGS.items():
        if f"<{tag}" in xml:
            problems.append(f"contains {label} ({tag})")
    # The rule used to be "no header or footer at all". That was the right rule when the
    # only reason to have one was contact details, which a parser that skips the footer
    # silently deletes. A name and a page number are different: both are already in the
    # body, so a parser dropping them costs nothing, and on a printed two-pager they are
    # what keeps page two attached to page one. So the check moved from where the content
    # sits to what the content is.
    if chrome:
        lost = _CONTACT_IN_CHROME.findall(chrome)
        if lost:
            problems.append(
                "contact details in a header or footer, which a parser may drop: "
                + ", ".join(sorted({m.lower() for m in lost})[:4])
            )
    if len(doc.tables) > 0:
        problems.append(f"{len(doc.tables)} table(s) via python-docx")
    for bad in PUNCTUATION_MAP:
        if bad in text:
            problems.append(f"contains U+{ord(bad):04X} {unicodedata.name(bad, '?')}")
    stray = sorted({c for c in text if ord(c) > 127} - set(PUNCTUATION_MAP))
    if stray:
        problems.append(
            "contains non-ASCII characters a parser may mangle: "
            + ", ".join(f"U+{ord(c):04X}" for c in stray[:8])
        )

    headings = [h for h in SECTION_HEADINGS.values() if h in text]

    return {
        "ok": not problems,
        "problems": problems,
        "headings_found": headings,
        "paragraphs": len(doc.paragraphs),
        "characters": len(text),
        "text": text,
    }


def keyword_coverage(path: Path, keywords: List[str]) -> Dict[str, bool]:
    """Which of the JD's must-have keywords actually appear in the rendered document."""
    text = audit(path)["text"].lower()
    return {kw: kw.lower() in text for kw in keywords}
